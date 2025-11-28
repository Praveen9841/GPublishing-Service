#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
GPublishing Services - HTML to Word Converter
Converts all HTML report files into a single Word document
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from bs4 import BeautifulSoup
    import os
    import re
    
    print("=" * 80)
    print("GPublishing Services - Report Converter (IMPROVED)")
    print("=" * 80)
    print("\nConverting HTML files to Word document...\n")
    
    # Create a new Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Configure heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
    
    # List of HTML files in order
    html_files = [
        'Report_00_TOC_Abstract.html',
        'Report_01_Chapter1_Introduction.html',
        'Report_02_Chapter2_SystemAnalysis.html',
        'Report_03_Chapter3_DevelopmentEnvironment.html',
        'Report_04_Chapter4_SystemDesign.html',
        'Report_05_Chapters5-9_References.html'
    ]
    
    # Process each HTML file
    for i, filename in enumerate(html_files, 1):
        print(f"[{i}/{len(html_files)}] Processing {filename}...")
        
        if not os.path.exists(filename):
            print(f"    WARNING: File not found - {filename}")
            continue
        
        # Read HTML file
        with open(filename, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        # Parse HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Extract body content
        body = soup.find('body')
        if not body:
            print(f"    WARNING: No body content found in {filename}")
            continue
        
        # Process each element recursively
        def process_element(elem, parent_list_type=None):
            if not elem or not hasattr(elem, 'name'):
                return
            
            if elem.name == 'h1':
                # Add heading 1
                text = elem.get_text().strip()
                if text:
                    p = doc.add_heading(text, level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            elif elem.name == 'h2':
                # Add heading 2
                text = elem.get_text().strip()
                if text:
                    doc.add_heading(text, level=2)
            
            elif elem.name == 'h3':
                # Add heading 3
                text = elem.get_text().strip()
                if text:
                    doc.add_heading(text, level=3)
            
            elif elem.name == 'p':
                # Add paragraph
                text = elem.get_text().strip()
                if text and len(text) > 1:
                    # Check if it's a strong/bold paragraph
                    if elem.find('strong'):
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    else:
                        p = doc.add_paragraph(text)
                        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            elif elem.name == 'ul':
                # Add unordered list
                for li in elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Bullet')
            
            elif elem.name == 'ol':
                # Add ordered list
                for li in elem.find_all('li', recursive=False):
                    text = li.get_text().strip()
                    if text:
                        doc.add_paragraph(text, style='List Number')
            
            elif elem.name == 'table':
                # Add table
                rows = elem.find_all('tr')
                if rows:
                    # Count max columns
                    max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
                    table = doc.add_table(rows=len(rows), cols=max_cols)
                    table.style = 'Light Grid Accent 1'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['th', 'td'])
                        for j, cell in enumerate(cells):
                            if j < max_cols:
                                cell_text = cell.get_text().strip()
                                table.rows[i].cells[j].text = cell_text
                                # Bold header cells
                                if cell.name == 'th':
                                    for paragraph in table.rows[i].cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.bold = True
                    
                    # Add spacing after table
                    doc.add_paragraph()
            
            elif elem.name == 'div':
                # Check for page break
                classes = elem.get('class', [])
                if classes and 'page-break' in classes:
                    doc.add_page_break()
                else:
                    # Process children of div
                    for child in elem.children:
                        if hasattr(child, 'name'):
                            process_element(child)
            
            elif elem.name == 'pre':
                # Add code block
                text = elem.get_text().strip()
                if text:
                    p = doc.add_paragraph(text)
                    p.style = 'Normal'
                    for run in p.runs:
                        run.font.name = 'Courier New'
                        run.font.size = Pt(10)
        
        for element in body.children:
            if hasattr(element, 'name'):
                process_element(element)
        
        print(f"    ✓ Completed {filename}")
    
    # Save the document
    output_filename = 'GPublishing_Services_Project_Report.docx'
    doc.save(output_filename)
    
    print("\n" + "=" * 80)
    print(f"✓ SUCCESS! Word document created: {output_filename}")
    print("=" * 80)
    print(f"\nFile location: {os.path.abspath(output_filename)}")
    print("\nYou can now open this file in Microsoft Word!")
    
except ImportError as e:
    print("\n" + "=" * 80)
    print("ERROR: Required libraries not installed")
    print("=" * 80)
    print("\nPlease install the required libraries by running:")
    print("\n  pip install python-docx beautifulsoup4")
    print("\nOr run:")
    print("\n  pip install python-docx beautifulsoup4 lxml")
    print("\nThen run this script again.")
    print("=" * 80)

except Exception as e:
    print("\n" + "=" * 80)
    print(f"ERROR: {str(e)}")
    print("=" * 80)
